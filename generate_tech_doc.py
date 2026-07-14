"""
SokoApp Technical Stack Documentation Generator
Run: python generate_tech_doc.py
Output: SokoApp_Technical_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ─── Brand colours ───────────────────────────────────────────────────
GREEN   = RGBColor(0x1D, 0x9E, 0x75)   # SokoApp green
DARK    = RGBColor(0x0F, 0x4C, 0x35)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GREY    = RGBColor(0x6B, 0x72, 0x80)
LIGHT   = RGBColor(0xF9, 0xFA, 0xFB)

HEX_GREEN  = '1D9E75'
HEX_DARK   = '0F4C35'
HEX_HEADER = '1D9E75'
HEX_LIGHT  = 'F0FDF4'
HEX_WHITE  = 'FFFFFF'


# ─── Helpers ─────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_col_widths(table, *widths_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = DARK
    # underline rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), HEX_GREEN)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = GREEN
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = GREY
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_stack_table(doc, rows, caption=None):
    """rows = list of (Category, Detail) tuples; first is treated as header."""
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(caption)
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = GREY

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # Header row
    hdr = table.add_row()
    for cell, txt in zip(hdr.cells, ['Technology / Layer', 'Details & Version']):
        shade_cell(cell, HEX_HEADER)
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(10)

    for cat, detail in rows:
        row = table.add_row()
        row.cells[0].text = cat
        row.cells[1].text = detail
        for i, cell in enumerate(row.cells):
            shade_cell(cell, HEX_LIGHT if i == 0 else HEX_WHITE)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

    set_col_widths(table, 2.3, 4.0)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ─── Cover page ──────────────────────────────────────────────────────

def build_cover(doc):
    # Top spacer
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('SokoApp')
    run.font.size  = Pt(42)
    run.font.bold  = True
    run.font.color.rgb = DARK

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Technical Architecture & Technology Stack')
    run.font.size  = Pt(18)
    run.font.color.rgb = GREEN

    doc.add_paragraph()
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule.add_run('─' * 48)
    r.font.color.rgb = GREEN

    doc.add_paragraph()
    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = desc_p.add_run(
        'A comprehensive multi-service platform covering mobile commerce,\n'
        'parcel delivery, savings groups, and financial administration.'
    )
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

    for _ in range(6):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run(f'Document generated: {datetime.date.today().strftime("%B %d, %Y")}')
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    add_page_break(doc)


# ─── 1. Executive Overview ───────────────────────────────────────────

def build_overview(doc):
    add_h1(doc, '1. Executive Overview')
    add_body(doc,
        'SokoApp is a full-stack, multi-tenant super-application designed for the Ghanaian '
        'market. It integrates grocery/product shopping (SokoShopper), parcel logistics '
        '(SokoDelivery), cooperative savings (SokoSusu), user identity (SokoAccount), '
        'and a web-based administration portal (SokoWeb) — all sharing a single unified '
        'Go backend API and a cross-platform React Native mobile client.'
    )

    add_h2(doc, 'Platform Components')
    components = [
        ('📱  Mobile App',       'React Native (Expo) — iOS, Android, Web'),
        ('🛒  SokoShopper',      'Grocery & product ordering module within the mobile app'),
        ('🚚  SokoDelivery',     'Parcel logistics — driver assignment, tracking, OTP delivery'),
        ('👤  SokoAccount',      'User identity, KYC verification, driver profiles, roles'),
        ('🤝  SokoSusu',         'Cooperative savings groups — contributions & rotating payouts'),
        ('🌐  SokoWeb',          'Flask-based superadmin portal for operations & finance'),
        ('⚙️   Go API Backend',  'Single unified REST API serving all mobile modules'),
    ]
    for icon_label, detail in components:
        add_bullet(doc, f'{icon_label}  —  {detail}')

    add_h2(doc, 'Architectural Principles')
    principles = [
        'Multi-database isolation: each domain has its own PostgreSQL database '
        '(sokoaccount, sokoshopper, sokodelivery, sokobank, sokosusu).',
        'Single Go binary: one process, one binary, multiple DB connections — no microservice overhead.',
        'Cross-DB joins at application layer: no foreign keys across databases; '
        'application code assembles data from multiple DB queries.',
        'Role-based access control (RBAC): JWT tokens carry role claims '
        '(user, driver, shopper_admin, superadmin) validated per route.',
        'Async task processing: Redis + Asynq for background jobs '
        '(push notifications, email, timeout watchers).',
        'Object storage: MinIO for all media assets (product images, KYC documents).',
        'Payment gateway: Paystack (Ghana) for all in-app payments.',
    ]
    for p in principles:
        add_bullet(doc, p)

    add_page_break(doc)


# ─── 2. Mobile App (SokoMobile) ──────────────────────────────────────

def build_mobile(doc):
    add_h1(doc, '2. Mobile App — SokoMobile (Frontend)')
    add_body(doc,
        'The mobile client is built with React Native via the Expo managed workflow. '
        'It is a single app that hosts all user-facing modules: shopping, parcel delivery, '
        'savings groups, and driver operations. The new architecture (Expo SDK 54, React 19) '
        'uses Expo Router for file-based navigation.'
    )

    add_stack_table(doc, [
        ('Core Framework',      'React Native 0.81.5 + React 19.1.0'),
        ('Development Platform','Expo SDK 54.0.35 (Managed Workflow)'),
        ('Language',            'TypeScript 5.9.2'),
        ('Navigation',          'Expo Router 6.0.23 — file-based routing (app/ directory)'),
        ('State Management',    'Zustand 5.0.13 — lightweight global store (authStore, etc.)'),
        ('HTTP Client',         'Axios 1.17.0 — API calls to Go backend'),
        ('Secure Storage',      'expo-secure-store 15.0.8 — JWT token persistence'),
        ('Async Storage',       '@react-native-async-storage/async-storage 2.2.0'),
        ('Maps & Location',     'react-native-maps 1.20.1 + expo-location 19.0.8'),
        ('Icons',               '@expo/vector-icons 15.0.3 (Ionicons, MaterialIcons)'),
        ('Image Picker',        'expo-image-picker 17.0.11 — KYC & avatar uploads'),
        ('Date/Time Picker',    '@react-native-community/datetimepicker 8.4.4'),
        ('WebView',             'react-native-webview 13.15.0 — Paystack payment gateway'),
        ('Web Browser',         'expo-web-browser 15.0.11 — OAuth / external links'),
        ('Deep Linking',        'expo-linking 8.0.12'),
        ('Safe Area',           'react-native-safe-area-context 5.6.0'),
        ('Screen Management',   'react-native-screens 4.16.0'),
        ('Constants',           'expo-constants 18.0.13'),
        ('Status Bar',          'expo-status-bar 3.0.9'),
        ('Architecture Flag',   'New Architecture enabled (Hermes engine, JSI bridge)'),
        ('Target Platforms',    'Android (edge-to-edge), iOS (tablet support), Web'),
    ])

    add_h2(doc, 'Key Mobile Modules')
    add_h3(doc, 'SokoShopper Module')
    add_bullet(doc, 'Store listing, category browsing, product detail')
    add_bullet(doc, 'Cart management with addons and quantity controls')
    add_bullet(doc, 'Checkout → Paystack WebView payment → OTP delivery confirmation')
    add_bullet(doc, 'Real-time order tracking (polling)')

    add_h3(doc, 'SokoDelivery Module (Customer side)')
    add_bullet(doc, 'Parcel booking form with pickup/dropoff address, description, vehicle type')
    add_bullet(doc, 'Price estimation before booking')
    add_bullet(doc, 'Live delivery status tracking')
    add_bullet(doc, 'Receiver confirmation via OTP')

    add_h3(doc, 'Driver Dashboard')
    add_bullet(doc, 'Active deliveries feed (parcel + shopper orders unified)')
    add_bullet(doc, 'Status flow: assigned → picked_up → in_transit → delivered (OTP)')
    add_bullet(doc, '5-minute pickup timer — cards lock if driver misses the window')
    add_bullet(doc, 'Live location sharing to backend via background GPS')
    add_bullet(doc, 'Earnings dashboard: total, weekly, available balance, cashout requests')
    add_bullet(doc, 'Cashout request modal (Mobile Money / Bank)')

    add_h3(doc, 'SokoSusu Module')
    add_bullet(doc, 'Group discovery, creation, join-request workflow')
    add_bullet(doc, 'Contribution history and rotating payout schedule')

    add_page_break(doc)


# ─── 3. SokoShopper (Backend domain) ─────────────────────────────────

def build_sokoshopper(doc):
    add_h1(doc, '3. SokoShopper — Backend Domain')
    add_body(doc,
        'SokoShopper is the backend domain within the Go API responsible for all '
        'product commerce: stores, products, cart, orders, and order delivery assignment. '
        'Its data lives exclusively in the sokoshopper PostgreSQL database.'
    )

    add_stack_table(doc, [
        ('Language',            'Go 1.23'),
        ('HTTP Framework',      'Gin 1.9.1'),
        ('ORM',                 'GORM 1.25.5 with gorm.io/driver/postgres 1.5.4'),
        ('Database',            'PostgreSQL (sokoshopper DB)'),
        ('Payment Gateway',     'Paystack Ghana — order payment initiation & webhook verification'),
        ('Object Storage',      'MinIO (minio-go v7.0.66) — product & store images'),
        ('Background Jobs',     'Asynq 0.24.1 on Redis — push notification tasks'),
        ('JWT Auth',            'golang-jwt/jwt v5.2.0'),
        ('UUID',                'google/uuid v1.6.0'),
        ('Env Config',          'godotenv 1.5.1'),
        ('Input Validation',    'go-playground/validator v10'),
    ])

    add_h2(doc, 'Database Models (sokoshopper)')
    models = [
        'Category — product categories',
        'Store — merchant stores (name, city, logo, status)',
        'StoreSection — sub-sections within a store',
        'Product — items with price, addons, images',
        'ProductAddon — optional extras per product',
        'Cart / CartItem — per-user cart persistence',
        'Order / OrderItem — placed orders with status state machine',
        'Payment — Paystack payment records',
        'OrderDelivery — driver assignment for shopper orders',
        'ShopCashoutRequest — admin-recorded store payout requests',
    ]
    for m in models:
        add_bullet(doc, m)

    add_h2(doc, 'Order Status Flow')
    flow = [
        'pending → payment_pending → payment_confirmed',
        'payment_confirmed → preparing → ready_for_pickup',
        'ready_for_pickup → assigned_to_driver → picked_up → in_transit → delivered',
        'Any state → cancelled  |  delivered → refunded',
    ]
    for f in flow:
        add_bullet(doc, f)

    add_h2(doc, 'Revenue Split')
    add_bullet(doc, '80% of order subtotal → store (store_payout)')
    add_bullet(doc, '20% of order subtotal → SokoApp platform')
    add_bullet(doc, '60% of delivery fee → driver (driver_earnings)')
    add_bullet(doc, '40% of delivery fee → SokoApp platform')

    add_page_break(doc)


# ─── 4. SokoDelivery ─────────────────────────────────────────────────

def build_sokodelivery(doc):
    add_h1(doc, '4. SokoDelivery — Backend Domain')
    add_body(doc,
        'SokoDelivery handles all parcel logistics: booking, driver broadcasting/assignment, '
        'GPS tracking, OTP-based delivery confirmation, and driver earnings calculation. '
        'Its data lives in the sokodelivery PostgreSQL database.'
    )

    add_stack_table(doc, [
        ('Language',            'Go 1.23'),
        ('HTTP Framework',      'Gin 1.9.1'),
        ('ORM',                 'GORM 1.25.5'),
        ('Database',            'PostgreSQL (sokodelivery DB)'),
        ('Payment',             'Paystack — parcel delivery fee payment'),
        ('Object Storage',      'MinIO — parcel photos if applicable'),
        ('Background Jobs',     'Asynq + Redis — pickup timeout watcher (5 min auto-reassign)'),
        ('Real-time Location',  'Driver POST /location endpoint — lat/lng stored per assignment'),
        ('UUID',                'google/uuid v1.6.0'),
        ('Crypto',              'golang.org/x/crypto — OTP hashing'),
    ])

    add_h2(doc, 'Database Models (sokodelivery)')
    models = [
        'DeliveryOrder — customer parcel booking (description, vehicle_type, pickup/dropoff coords & address)',
        'DeliveryAssignment — driver-to-order link with status, earnings split, delivered_at',
        'DeliveryZone — geographic pricing zones',
        'DriverEarning — per-delivery earning record (amount, platform_cut)',
        'DriverCashoutRequest — driver payout requests (pending / paid / rejected)',
    ]
    for m in models:
        add_bullet(doc, m)

    add_h2(doc, 'Delivery Status Flow')
    flow = [
        'broadcast → accepted (driver accepts broadcast)',
        'accepted → picked_up (driver marks pickup)',
        'picked_up → in_transit',
        'in_transit → delivered (OTP required from receiver)',
    ]
    for f in flow:
        add_bullet(doc, f)

    add_h2(doc, 'Earnings Split (Parcels)')
    add_bullet(doc, '60% of delivery fee → driver (stored in DeliveryAssignment.driver_earnings)')
    add_bullet(doc, '40% of delivery fee → SokoApp (stored in DeliveryAssignment.platform_cut)')
    add_bullet(doc, 'DriverEarning record created per completed delivery for audit trail')

    add_h2(doc, 'Background Worker — Pickup Timeout Watcher')
    add_bullet(doc, 'Runs every 30 seconds via robfig/cron v3')
    add_bullet(doc, 'Finds assignments not marked picked_up within 5 minutes of acceptance')
    add_bullet(doc, 'Auto-resets status to broadcast for re-assignment by admin')

    add_page_break(doc)


# ─── 5. SokoAccount ──────────────────────────────────────────────────

def build_sokoaccount(doc):
    add_h1(doc, '5. SokoAccount — Identity & Auth Domain')
    add_body(doc,
        'SokoAccount is the central identity service. It manages all user profiles, '
        'authentication (JWT), role assignments, KYC document submission, and driver '
        'profile management. Every other service resolves user identity by querying '
        'this database at the application layer.'
    )

    add_stack_table(doc, [
        ('Language',            'Go 1.23'),
        ('HTTP Framework',      'Gin 1.9.1'),
        ('ORM',                 'GORM 1.25.5'),
        ('Database',            'PostgreSQL (sokoaccount DB)'),
        ('Auth',                'JWT (golang-jwt/jwt v5.2.0) — access + refresh token pair'),
        ('Password Hashing',    'bcrypt via golang.org/x/crypto'),
        ('Object Storage',      'MinIO — KYC document uploads, profile avatars'),
        ('Email / Notifications','Asynq task queue → background email dispatch'),
        ('OTP Verification',    'In-app OTP for phone/email verification (OTPVerification model)'),
        ('UUID',                'google/uuid v1.6.0'),
        ('Env Config',          'godotenv 1.5.1'),
    ])

    add_h2(doc, 'Database Models (sokoaccount)')
    models = [
        'User — core identity (full_name, email, phone, avatar, status)',
        'UserRole — many-to-many role assignments (user, driver, shopper_admin, superadmin)',
        'RefreshToken — secure token storage with expiry',
        'OTPVerification — phone/email OTP records',
        'DriverProfile — vehicle details, license, bank/MoMo info',
        'KYCSubmission / KYCDocument — identity document workflow',
        'AuditLog — action log for sensitive operations',
    ]
    for m in models:
        add_bullet(doc, m)

    add_h2(doc, 'Auth Flow')
    add_bullet(doc, 'Register → optional email OTP verify → login')
    add_bullet(doc, 'Login returns access_token (short TTL) + refresh_token (long TTL)')
    add_bullet(doc, 'All protected routes use JWTAuthMiddleware + AuthorizeRoles RBAC middleware')
    add_bullet(doc, 'Refresh endpoint issues new access token from valid refresh token')
    add_bullet(doc, 'Driver application creates DriverProfile + pending KYC → superadmin approves')

    add_page_break(doc)


# ─── 6. SokoSusu ─────────────────────────────────────────────────────

def build_sokosusu(doc):
    add_h1(doc, '6. SokoSusu — Cooperative Savings Domain')
    add_body(doc,
        'SokoSusu implements the traditional West African "susu" rotating savings model. '
        'Users form groups, contribute on a schedule, and receive rotating lump-sum payouts. '
        'Its data lives in the sokosusu PostgreSQL database.'
    )

    add_stack_table(doc, [
        ('Language',            'Go 1.23'),
        ('HTTP Framework',      'Gin 1.9.1'),
        ('ORM',                 'GORM 1.25.5'),
        ('Database',            'PostgreSQL (sokosusu DB)'),
        ('UUID',                'google/uuid v1.6.0'),
        ('Auth',                'JWT — same token issued by SokoAccount'),
        ('Background Jobs',     'Asynq — payout notifications'),
    ])

    add_h2(doc, 'Database Models (sokosusu)')
    models = [
        'SusuGroup — group metadata (name, description, contribution amount, frequency, cycle)',
        'SusuMember — member-to-group link with join date and payout order',
        'SusuJoinRequest — pending/approved/rejected join requests',
        'SusuContribution — individual contribution payment records',
        'SusuPayout — rotating payout records per member',
    ]
    for m in models:
        add_bullet(doc, m)

    add_h2(doc, 'Key Features')
    add_bullet(doc, 'Group discovery — browse public groups, filter by contribution amount')
    add_bullet(doc, 'Join request workflow — group admin approves or rejects applicants')
    add_bullet(doc, 'Contribution tracking — per-member contribution history')
    add_bullet(doc, 'Rotating payout — each cycle one member receives the pool')
    add_bullet(doc, 'Admin controls — create group, manage members, record payouts')

    add_page_break(doc)


# ─── 7. SokoWeb ──────────────────────────────────────────────────────

def build_sokoweb(doc):
    add_h1(doc, '7. SokoWeb — Superadmin Web Portal')
    add_body(doc,
        'SokoWeb is a server-rendered web application used by SokoApp superadmins. '
        'It connects to all five PostgreSQL databases simultaneously to provide a unified '
        'operations dashboard — managing users, drivers, stores, deliveries, and finances.'
    )

    add_stack_table(doc, [
        ('Language',            'Python 3.x'),
        ('Web Framework',       'Flask 3.0.3'),
        ('ORM',                 'Flask-SQLAlchemy 3.1.1 + SQLAlchemy 2.0.36'),
        ('Database Driver',     'psycopg2-binary — PostgreSQL adapter'),
        ('DB Architecture',     'Multi-bind SQLAlchemy — 5 separate DB connections via SQLALCHEMY_BINDS'),
        ('Auth',                'Session-based login + PyJWT 2.9.0 for generating backend tokens'),
        ('Templating',          'Jinja2 (bundled with Flask) — server-side HTML rendering'),
        ('Frontend CSS',        'Bootstrap 5.3 + Bootstrap Icons'),
        ('Image Processing',    'Pillow 10.4.0'),
        ('Push Notifications',  'exponent-server-sdk 2.0.0 — Expo push notification dispatch'),
        ('Realtime (optional)', 'python-socketio 5.11.1 + python-engineio 4.9.1'),
        ('HTTP Requests',       'requests 2.34.2 — calls to Go API backend'),
        ('Env Config',          'python-dotenv 1.0.1'),
        ('Email Validation',    'email-validator 2.1.0'),
        ('Word Export',         'python-docx 1.1.2 — this document generation'),
    ])

    add_h2(doc, 'Web Portal Modules')
    modules = [
        ('Dashboard',     'KPI overview — orders, drivers, revenue at a glance'),
        ('Shopper Admin', 'Store management, product catalogue, order list & status updates'),
        ('Account Admin', 'User list, KYC review, driver approval & suspension'),
        ('Delivery Admin','Parcel order board, driver assignment, delivery map'),
        ('Finance',       'Track Expenses — full revenue breakdown, pending cashouts, Mark Paid'),
        ('Susu Admin',    'Savings group oversight, contribution & payout records'),
    ]
    for name, desc in modules:
        add_bullet(doc, f'{name}  —  {desc}')

    add_h2(doc, 'Finance Dashboard Features')
    add_bullet(doc, 'Grand total SokoApp earnings (20% product + 40% delivery fees)')
    add_bullet(doc, 'Per-driver earnings: parcel + shopper, paid out, pending cashout, available balance')
    add_bullet(doc, 'Per-store breakdown: gross sales, 80% store payout, SokoApp 20% cut')
    add_bullet(doc, 'Pending driver cashout queue with one-click "Mark Paid"')
    add_bullet(doc, 'Admin-recorded store payout modal with "Mark Paid" workflow')
    add_bullet(doc, 'Date range filtering across all completed deliveries and orders')
    add_bullet(doc, 'Printable report view')

    add_page_break(doc)


# ─── 8. Backend API Architecture ─────────────────────────────────────

def build_api(doc):
    add_h1(doc, '8. Go API Backend — Shared Services')
    add_body(doc,
        'All five domains are served by a single Go 1.23 binary exposing a versioned '
        'REST API at /api/v1/. The server is structured as a monolith with clear internal '
        'package separation — not microservices — which simplifies deployment while still '
        'enforcing domain isolation at the database layer.'
    )

    add_stack_table(doc, [
        ('Language',            'Go 1.23'),
        ('HTTP Router',         'Gin 1.9.1 — high-performance HTTP framework'),
        ('CORS',                'gin-contrib/cors 1.4.0'),
        ('ORM',                 'GORM 1.25.5 — all five DB connections'),
        ('DB Driver',           'pgx v5.4.3 — PostgreSQL driver'),
        ('Authentication',      'golang-jwt/jwt v5.2.0 — HS256 signed JWTs'),
        ('Password Hashing',    'bcrypt via golang.org/x/crypto'),
        ('Task Queue (broker)', 'Redis via redis/go-redis v9.0.5'),
        ('Task Queue (worker)', 'Asynq 0.24.1 — background job processing'),
        ('Cron Jobs',           'robfig/cron v3.0.1 — scheduled pickup timeout watcher'),
        ('Object Storage',      'MinIO minio-go v7.0.66 — S3-compatible media storage'),
        ('Payment Gateway',     'Paystack — webhook + REST verification'),
        ('UUID Generation',     'google/uuid v1.6.0'),
        ('JSON Performance',    'bytedance/sonic (fast JSON encoder/decoder via Gin)'),
        ('Environment Config',  'joho/godotenv v1.5.1'),
        ('Input Validation',    'go-playground/validator v10.14.0'),
        ('Logging',             'sirupsen/logrus v1.9.3'),
    ])

    add_h2(doc, 'API Route Groups')
    routes = [
        '/api/v1/auth/*         — register, login, refresh, logout, /me',
        '/api/v1/shopper/*      — stores, products, cart, orders, payment',
        '/api/v1/delivery/*     — parcel booking, tracking, price estimate',
        '/api/v1/delivery/driver/* — driver deliveries, status updates, OTP, location, earnings, cashout',
        '/api/v1/account/*      — KYC submission, driver application',
        '/api/v1/susu/*         — groups, members, contributions, payouts',
        '/api/v1/media/*        — MinIO media upload & serve proxy',
        '/api/v1/webhooks/*     — Paystack payment webhook handler',
    ]
    for r in routes:
        add_bullet(doc, r)

    add_h2(doc, 'Security Middleware')
    add_bullet(doc, 'JWTAuthMiddleware — validates Bearer token on every protected route')
    add_bullet(doc, 'AuthorizeRoles — checks role claim against allowed roles per route group')
    add_bullet(doc, 'CORS — allows all origins in dev (configurable for production whitelist)')
    add_bullet(doc, 'Trusted proxies set to nil — prevents IP spoofing via proxy headers')

    add_page_break(doc)


# ─── 9. Database Architecture ────────────────────────────────────────

def build_databases(doc):
    add_h1(doc, '9. Database Architecture')
    add_body(doc,
        'SokoApp uses five isolated PostgreSQL databases. Each domain owns its schema '
        'completely — there are no cross-database foreign keys. Joins across domains are '
        'performed at the application layer by querying multiple databases and merging '
        'results in Go or Python code.'
    )

    add_stack_table(doc, [
        ('sokoaccount',     'Users, roles, refresh tokens, OTP, driver profiles, KYC documents, audit log'),
        ('sokoshopper',     'Categories, stores, products, carts, orders, payments, shop cashout requests'),
        ('sokodelivery',    'Parcel orders, driver assignments, delivery zones, driver earnings, cashout requests'),
        ('sokobank',        'Bank accounts, transactions, transfers, account statements'),
        ('sokosusu',        'Susu groups, members, join requests, contributions, payouts'),
    ])

    add_h2(doc, 'Cross-DB Join Pattern')
    add_body(doc,
        'Example: Finance dashboard needs driver name (sokoaccount) + delivery earnings '
        '(sokodelivery) + shopper order delivery (sokoshopper):'
    )
    add_bullet(doc, 'Step 1 — Query sokodelivery for all completed DeliveryAssignments → get driver_id list')
    add_bullet(doc, 'Step 2 — Query sokoaccount.users WHERE id IN (driver_id list) → get names')
    add_bullet(doc, 'Step 3 — Query sokoshopper.order_deliveries WHERE driver_id IN (...) → get shopper deliveries')
    add_bullet(doc, 'Step 4 — Merge all three result sets in Python/Go using dictionaries keyed by user ID')

    add_h2(doc, 'Schema Management')
    add_bullet(doc, 'Go backend: GORM AutoMigrate on every startup — creates/alters tables, never drops')
    add_bullet(doc, 'Flask/SokoWeb: SQLAlchemy db.create_all() on startup — creates missing tables')
    add_bullet(doc, 'PostgreSQL enum types created manually before AutoMigrate (GORM limitation)')
    add_bullet(doc, 'UUID primary keys throughout (uuid_generate_v4()) — safe for distributed use')

    add_page_break(doc)


# ─── 10. Infrastructure & Deployment ─────────────────────────────────

def build_infra(doc):
    add_h1(doc, '10. Infrastructure & Deployment')

    add_stack_table(doc, [
        ('Go API Server',       'Single binary, port 8082 (configurable via PORT env var)'),
        ('Flask Web Portal',    'Python WSGI, typically port 5000 (Gunicorn in production)'),
        ('PostgreSQL',          '5 databases on a single PostgreSQL instance (or separate instances)'),
        ('Redis',               'Single Redis instance — Asynq job queue + caching'),
        ('MinIO',               'Self-hosted S3-compatible object storage — media & documents'),
        ('Paystack',            'External SaaS — payments, webhooks, refunds (Ghana)'),
        ('Expo Push',           'Expo Notifications service — mobile push via exponent-server-sdk'),
        ('Environment Config',  '.env files loaded by godotenv (Go) and python-dotenv (Flask)'),
        ('Platform',            'Windows 11 Pro (development) — deployable to Linux VPS / Docker'),
    ])

    add_h2(doc, 'Key Environment Variables')
    env_vars = [
        'DB_HOST, DB_PORT, DB_USER, DB_PASSWORD — PostgreSQL connection',
        'SOKOACCOUNT_DB, SOKOSHOPPER_DB, SOKODELIVERY_DB, SOKOBANK_DB, SOKOSUSU_DB',
        'REDIS_HOST, REDIS_PORT — Asynq task queue',
        'JWT_SECRET — shared HS256 signing key (Go + Flask both use this)',
        'PAYSTACK_SECRET_KEY — Paystack Ghana API secret',
        'MINIO_ENDPOINT, MINIO_ACCESS_KEY, MINIO_SECRET_KEY, MINIO_BUCKET',
        'API_BASE_URL — Go backend URL used by Flask to proxy media',
        'PORT — Go API listen port (default 8082)',
    ]
    for v in env_vars:
        add_bullet(doc, v)

    add_page_break(doc)


# ─── 11. Overall Tech Summary ─────────────────────────────────────────

def build_summary(doc):
    add_h1(doc, '11. Overall Technology Summary')

    add_stack_table(doc, [
        ('Category',                 'Technology'),  # will be styled as header by add_stack_table
    ])

    # Overwrite — just use a plain wider table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0]
    for cell, txt in zip(hdr.cells, ['Layer', 'Technology', 'Version / Notes']):
        shade_cell(cell, HEX_DARK)
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.bold  = True
        r.font.color.rgb = WHITE
        r.font.size  = Pt(10)

    rows_data = [
        ('Mobile Language',     'TypeScript',               '5.9.2'),
        ('Mobile Framework',    'React Native',             '0.81.5 (Expo SDK 54)'),
        ('Mobile Navigation',   'Expo Router',              '6.0.23 — file-based'),
        ('Mobile State',        'Zustand',                  '5.0.13'),
        ('Mobile HTTP',         'Axios',                    '1.17.0'),
        ('Mobile Maps',         'react-native-maps',        '1.20.1'),
        ('Backend Language',    'Go',                       '1.23'),
        ('Backend Framework',   'Gin',                      '1.9.1'),
        ('Backend ORM',         'GORM',                     '1.25.5'),
        ('Databases',           'PostgreSQL',               '5 isolated DBs'),
        ('Task Queue',          'Redis + Asynq',            'v9.0.5 + 0.24.1'),
        ('Object Storage',      'MinIO',                    'minio-go v7'),
        ('Authentication',      'JWT (HS256)',               'golang-jwt v5'),
        ('Payments',            'Paystack',                  'Ghana payment gateway'),
        ('Admin Language',      'Python 3.x',               'Flask 3.0.3'),
        ('Admin ORM',           'SQLAlchemy',               '2.0.36 (multi-bind)'),
        ('Admin Templates',     'Jinja2 + Bootstrap 5',     'Server-side rendered'),
        ('Admin Auth',          'Flask sessions + PyJWT',   '2.9.0'),
        ('Scheduled Jobs',      'robfig/cron',              'v3.0.1'),
        ('Media Serving',       'MinIO + Gin proxy',        'S3-compatible'),
        ('Mobile Storage',      'expo-secure-store',        '15.0.8'),
        ('Push Notifications',  'Expo Notifications',       'exponent-server-sdk 2.0'),
    ]

    for layer, tech, ver in rows_data:
        row = table.add_row()
        row.cells[0].text = layer
        row.cells[1].text = tech
        row.cells[2].text = ver
        shade_cell(row.cells[0], HEX_LIGHT)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True

    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(2.2)
        row.cells[2].width = Inches(2.3)

    doc.add_paragraph()

    add_h2(doc, 'Glossary')
    terms = [
        ('RBAC',    'Role-Based Access Control — routes are gated by user role claims in JWT'),
        ('GORM',    'Go ORM library — handles all SQL generation and schema migration'),
        ('Asynq',   'Go async task queue backed by Redis — used for background jobs'),
        ('MinIO',   'Self-hosted S3-compatible object store — holds all uploaded media'),
        ('Paystack','Ghanaian payment gateway — handles card, mobile money, bank payments'),
        ('Susu',    'Traditional West African cooperative savings group — rotating payouts'),
        ('KYC',     'Know Your Customer — identity document verification for drivers'),
        ('OTP',     'One-Time Password — used to confirm delivery and phone verification'),
        ('JWT',     'JSON Web Token — stateless auth token signed with shared HS256 secret'),
    ]
    table2 = doc.add_table(rows=0, cols=2)
    table2.style = 'Table Grid'
    hdr2 = table2.add_row()
    for cell, txt in zip(hdr2.cells, ['Term', 'Definition']):
        shade_cell(cell, HEX_GREEN)
        r = cell.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for term, defn in terms:
        row = table2.add_row()
        row.cells[0].text = term
        row.cells[1].text = defn
        shade_cell(row.cells[0], HEX_LIGHT)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
    set_col_widths(table2, 1.2, 5.1)


# ─── Build & save ─────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin   = Inches(1)
    section.bottom_margin= Inches(1)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    build_cover(doc)
    build_overview(doc)
    build_mobile(doc)
    build_sokoshopper(doc)
    build_sokodelivery(doc)
    build_sokoaccount(doc)
    build_sokosusu(doc)
    build_sokoweb(doc)
    build_api(doc)
    build_databases(doc)
    build_infra(doc)
    build_summary(doc)

    out = 'SokoApp_Technical_Documentation.docx'
    doc.save(out)
    print(f'✅  Document saved → {os.path.abspath(out)}')
    return out


if __name__ == '__main__':
    build_document()
